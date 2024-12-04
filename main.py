import tkinter as tk
from tkinter import ttk, filedialog
from database import Database  # Предполагаем, что у вас есть класс для работы с базой данных
from docx import Document

class TaxCalculatorApp:
    def __init__(self, root, db):
        self.root = root
        self.db = db
        self.root.title("Налоговый калькулятор")
        self.root.geometry("400x500")  # Set a default window size
        self.root.resizable(True, True)  # Allow resizing

        self.deductions = self.load_deductions()  # Загрузка вычетов из базы данных

        # Main Frame for the content
        self.main_frame = ttk.Frame(self.root, padding="20")
        self.main_frame.grid(row=0, column=0, sticky="nsew")
        self.main_frame.columnconfigure(0, weight=1)  # Make the first column expand
        self.main_frame.rowconfigure(0, weight=1)  # Make the first row expand

        # Общая сумма, включая НДФЛ
        self.sum_label = ttk.Label(self.main_frame, text="Общая сумма, включая НДФЛ:", anchor="w")
        self.sum_label.grid(row=0, column=0, sticky="w", pady=5)
        self.sum_entry = ttk.Entry(self.main_frame, width=20)
        self.sum_entry.grid(row=0, column=1, pady=5)
        self.sum_entry.insert(0, "100")  # Default value

        # Ставка налога
        self.tax_rate_label = ttk.Label(self.main_frame, text="Ставка налога:", anchor="w")
        self.tax_rate_label.grid(row=1, column=0, sticky="w", pady=5)
        self.tax_rate_combobox = ttk.Combobox(self.main_frame, values=["6%", "13%", "15%"], width=10)
        self.tax_rate_combobox.grid(row=1, column=1, pady=5)
        self.tax_rate_combobox.set("13%")  # Default value

        # Основные вычеты (deductions)
        self.deductions_frame = ttk.LabelFrame(self.main_frame, text="Основные вычеты", padding="10")
        self.deductions_frame.grid(row=2, column=0, columnspan=2, pady=10, sticky="ew")

        self.deduction_vars = {}  # Словарь для хранения состояния чекбоксов

        # Создаем чекбоксы для каждого вычета
        row = 0
        for deduction, amount in self.deductions.items():
            var = tk.BooleanVar()
            checkbox = ttk.Checkbutton(self.deductions_frame, text=f"{deduction} - {amount} Р", variable=var)
            checkbox.grid(row=row, column=0, sticky="w", padx=5)
            self.deduction_vars[deduction] = var
            row += 1

        # Кнопка расчета
        self.calculate_button = ttk.Button(self.main_frame, text="Рассчитать", command=self.calculate_tax)
        self.calculate_button.grid(row=3, column=0, columnspan=2, pady=15)

        # Результаты
        self.result_label = ttk.Label(self.main_frame, text="Сумма за вычетом НДФЛ:", anchor="w")
        self.result_label.grid(row=4, column=0, sticky="w", pady=5)
        self.result_value_label = ttk.Label(self.main_frame, text="0 ₽", font=("Arial", 14, "bold"))
        self.result_value_label.grid(row=4, column=1, sticky="e", pady=5)

        # Налог
        self.tax_amount_label = ttk.Label(self.main_frame, text="Налог:", anchor="w")
        self.tax_amount_label.grid(row=5, column=0, sticky="w", pady=5)
        self.tax_amount_value_label = ttk.Label(self.main_frame, text="0 ₽", font=("Arial", 14, "bold"))
        self.tax_amount_value_label.grid(row=5, column=1, sticky="e", pady=5)

        # Описание расчетов
        self.description_label = ttk.Label(self.main_frame, text="Расчет НДФЛ", anchor="w", font=("Arial", 12, "bold"))
        self.description_label.grid(row=6, column=0, columnspan=2, pady=10)
        self.explanation_label = ttk.Label(self.main_frame,
                                           text="НДФЛ = общая сумма, включая НДФЛ x ставка налога\nСумма за вычетом НДФЛ = общая сумма - НДФЛ",
                                           justify="left", font=("Arial", 10))
        self.explanation_label.grid(row=7, column=0, columnspan=2, pady=5)

        # Место для кнопок "Экспорт в Word" (они будут появляться после расчета)
        self.export_button_frame = ttk.Frame(self.main_frame)
        self.export_button_frame.grid(row=8, column=0, columnspan=2, pady=10)

    def load_deductions(self):
        query = "SELECT name, amount FROM deductions"  # Запрос для получения всех вычетов
        result = self.db.fetch_all(query)  # Получаем все вычеты из базы

        # Преобразуем результат в словарь, где ключ — это название вычета, а значение — сумма
        return {deduction[0]: deduction[1] for deduction in result}

    def calculate_tax(self):
        # Получаем входные значения
        try:
            total_sum = float(self.sum_entry.get())  # Общая сумма
        except ValueError:
            self.result_value_label.config(text="Неверный ввод!")
            self.tax_amount_value_label.config(text="Неверный ввод!")
            return

        tax_rate = self.tax_rate_combobox.get()

        # Проверяем, что ставка налога не пустая
        if not tax_rate:
            self.result_value_label.config(text="Выберите ставку налога!")
            return

        try:
            tax_rate_percent = int(tax_rate[:-1]) / 100  # Убираем '%' и конвертируем в число
        except ValueError:
            self.result_value_label.config(text="Неверный формат ставки налога!")
            return

        # Вычеты
        deductions = 0
        for deduction, var in self.deduction_vars.items():
            if var.get():
                deductions += float(self.deductions[deduction])  # Приводим к float

        # Рассчитываем НДФЛ (налог)
        tax_amount = total_sum * tax_rate_percent  # НДФЛ = общая сумма * ставка налога
        tax_included = total_sum - tax_amount  # Сумма без налога

        # Финальная сумма после вычета НДФЛ и вычетов
        result_sum = tax_included - deductions
        if result_sum < 0:  # Если вычеты больше суммы, делаем сумму 0
            result_sum = 0

        # Обновляем метки с результатами
        self.result_value_label.config(text=f"{result_sum:.2f} ₽")
        self.tax_amount_value_label.config(text=f"{tax_amount:.2f} ₽")

        # Обновление описания расчетов
        self.explanation_label.config(
            text=f"НДФЛ = {total_sum:.2f} ₽ x {tax_rate_percent * 100}% = {tax_amount:.2f} ₽\n"
                 f"Сумма за вычетом НДФЛ = {tax_included:.2f} ₽ - {deductions:.2f} ₽ = {result_sum:.2f} ₽"
        )

        # Показать кнопки "Экспорт в Word"
        self.show_export_buttons()

    def show_export_buttons(self):
        # Удалить старые кнопки, если они есть
        for widget in self.export_button_frame.winfo_children():
            widget.destroy()

        # Кнопка "Экспорт в Word"
        export_button = ttk.Button(self.export_button_frame, text="Экспорт в Word", command=self.export_to_word)
        export_button.grid(row=0, column=0, padx=10)

        # Кнопка "Отмена"
        cancel_button = ttk.Button(self.export_button_frame, text="Отмена", command=self.cancel_export)
        cancel_button.grid(row=0, column=1, padx=10)

    def export_to_word(self):
        # Создаем новый документ
        doc = Document()

        # Добавляем информацию в документ
        doc.add_heading('Расчет НДФЛ', 0)

        total_sum = float(self.sum_entry.get())
        tax_rate = self.tax_rate_combobox.get()
        tax_rate_percent = int(tax_rate[:-1]) / 100  # Убираем '%' и конвертируем в число

        deductions = 0
        for deduction, var in self.deduction_vars.items():
            if var.get():
                deductions += float(self.deductions[deduction])  # Приводим к float

        # Рассчитываем НДФЛ
        tax_amount = total_sum * tax_rate_percent
        tax_included = total_sum - tax_amount  # Сумма без налога

        # Обновление текста для документа
        doc.add_paragraph(f"Общая сумма, включая НДФЛ: {total_sum:.2f} ₽")
        doc.add_paragraph(f"Ставка налога: {tax_rate_percent * 100}%")
        doc.add_paragraph(f"Налог (НДФЛ): {tax_amount:.2f} ₽")
        doc.add_paragraph(f"Вычеты: {deductions:.2f} ₽")
        doc.add_paragraph(f"Сумма за вычетом НДФЛ: {tax_included - deductions:.2f} ₽")

        # Сохранение документа
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if file_path:
            doc.save(file_path)

    def cancel_export(self):
        # Удаление кнопок экспорта
        for widget in self.export_button_frame.winfo_children():
            widget.destroy()

if __name__ == "__main__":
    db = Database()  # Подключаемся к базе данных
    root = tk.Tk()
    app = TaxCalculatorApp(root, db)
    root.mainloop()
